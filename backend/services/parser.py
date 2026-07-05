"""
Phase 1 — Document Parser
Supports: PDF, DOCX, TXT, CSV, JSON
Returns the raw extracted text from any supported file type.
"""
import json
import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def parse_document(file_path: Path, file_name: str) -> str:
    """
    Extract raw text from a file based on its extension.
    Returns a single string of the document's text content.
    """
    ext = Path(file_name).suffix.lower()

    try:
        if ext == ".pdf":
            return _parse_pdf(file_path)
        elif ext == ".docx":
            return _parse_docx(file_path)
        elif ext == ".txt":
            return _parse_txt(file_path)
        elif ext == ".csv":
            return _parse_csv(file_path)
        elif ext == ".json":
            return _parse_json(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    except Exception as e:
        logger.error(f"Failed to parse document '{file_name}': {e}")
        raise


def _parse_pdf(file_path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(file_path))
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            pages.append(f"[Page {i + 1}]\n{text.strip()}")
    
    if not pages:
        raise ValueError("PDF appears to be empty or contains only images.")
    
    return "\n\n".join(pages)


def _parse_docx(file_path: Path) -> str:
    from docx import Document

    doc = Document(str(file_path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                paragraphs.append(row_text)
    
    if not paragraphs:
        raise ValueError("DOCX appears to be empty.")
    
    return "\n\n".join(paragraphs)


def _parse_txt(file_path: Path) -> str:
    # Try utf-8, fall back to latin-1
    try:
        text = file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = file_path.read_text(encoding="latin-1")
    
    if not text.strip():
        raise ValueError("TXT file is empty.")
    
    return text.strip()


def _parse_csv(file_path: Path) -> str:
    import pandas as pd

    df = pd.read_csv(str(file_path))
    if df.empty:
        raise ValueError("CSV file is empty.")
    
    # Return both the headers and row-by-row text
    lines = [",".join(str(c) for c in df.columns)]
    for _, row in df.iterrows():
        lines.append(",".join(str(v) for v in row.values))
    
    return "\n".join(lines)


def _parse_json(file_path: Path) -> str:
    try:
        text = file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = file_path.read_text(encoding="latin-1")
    
    data = json.loads(text)
    
    # Pretty-print JSON to preserve structure for chunking
    return json.dumps(data, indent=2, ensure_ascii=False)
